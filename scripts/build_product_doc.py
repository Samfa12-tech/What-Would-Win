from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
MARKDOWN = ROOT / 'docs' / 'What_Would_Win_Product_Plan.md'
OUT = ROOT / 'docs' / 'What_Would_Win_Product_Plan.docx'
ASSETS = ROOT / 'docs' / 'assets'

NAVY = '0B1220'
NAVY_2 = '142033'
GOLD = 'C9A86A'
GOLD_LIGHT = 'EAD7A5'
INK = '1F2937'
MUTED = '596579'
PALE = 'F4F1E8'
PALE_BLUE = 'EEF2F7'
WHITE = 'FFFFFF'
BORDER = 'CAD1DB'


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in('w:tcBorders')
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = 'w:{}'.format(edge)
            element = tc_borders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tc_borders.append(element)
            for key in ['sz', 'val', 'color', 'space']:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def set_cell_margins(cell, top=100, start=100, bottom=100, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run('Page ')
    run.font.size = Pt(8)
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'PAGE'
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_run_font(run, name='Arial', size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_inline(paragraph, text: str, base_size=9.2, base_color=INK, code=False) -> None:
    # Supports bold, italics, inline code and raw URLs. Deliberately conservative.
    token_re = re.compile(r'(\*\*.+?\*\*|`.+?`|(?<!\*)\*[^*]+?\*(?!\*)|https?://\S+)')
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, 'Courier New' if code else 'Arial', base_size, base_color)
        token = match.group(0)
        if token.startswith('**'):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, 'Arial', base_size, base_color, bold=True)
        elif token.startswith('`'):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, 'Courier New', base_size - 0.2, NAVY)
            run.font.highlight_color = None
        elif token.startswith('*'):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, 'Arial', base_size, base_color, italic=True)
        else:
            run = paragraph.add_run(token.rstrip('.,);'))
            set_run_font(run, 'Arial', base_size - 0.4, '1B4D77')
            run.underline = True
            suffix = token[len(token.rstrip('.,);')):]
            if suffix:
                tail = paragraph.add_run(suffix)
                set_run_font(tail, 'Arial', base_size, base_color)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, 'Courier New' if code else 'Arial', base_size, base_color)


def add_styled_paragraph(doc, text: str, *, style=None, align=None, spacing_after=4, spacing_before=0, keep=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.space_before = Pt(spacing_before)
    p.paragraph_format.line_spacing = 1.08
    p.paragraph_format.keep_together = keep
    add_inline(p, text)
    return p


def add_callout(doc, text: str, fill=PALE, border=GOLD):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, left={'val':'single','sz':'18','color':border,'space':'0'},
                    top={'val':'single','sz':'4','color':border,'space':'0'},
                    bottom={'val':'single','sz':'4','color':border,'space':'0'},
                    right={'val':'single','sz':'4','color':border,'space':'0'})
    set_cell_margins(cell, 140, 180, 140, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, text, base_size=9.1)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_markdown_table(doc, rows: list[list[str]]):
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=cols)
    table.autofit = True
    for r_idx, row_data in enumerate(rows):
        cells = table.add_row().cells
        for c_idx in range(cols):
            value = row_data[c_idx].strip() if c_idx < len(row_data) else ''
            cell = cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, 85, 105, 85, 105)
            set_cell_border(cell,
                top={'val':'single','sz':'4','color':BORDER,'space':'0'},
                bottom={'val':'single','sz':'4','color':BORDER,'space':'0'},
                left={'val':'single','sz':'4','color':BORDER,'space':'0'},
                right={'val':'single','sz':'4','color':BORDER,'space':'0'})
            if r_idx == 0:
                set_cell_shading(cell, NAVY_2)
            elif r_idx % 2 == 0:
                set_cell_shading(cell, 'F7F9FB')
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline(p, value, base_size=7.8, base_color=WHITE if r_idx == 0 else INK)
            for run in p.runs:
                if r_idx == 0:
                    run.bold = True
        if r_idx == 0:
            set_repeat_table_header(table.rows[0])
    table.rows[0].height = Cm(0.65)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def parse_table(lines: list[str], start: int):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith('|'):
        parts = [part.strip() for part in lines[i].strip().strip('|').split('|')]
        rows.append(parts)
        i += 1
    if len(rows) >= 2 and all(re.fullmatch(r':?-{3,}:?', p.replace(' ', '')) for p in rows[1]):
        rows.pop(1)
    return rows, i


def add_code_block(doc, code: str):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0,0)
    set_cell_shading(cell, 'F1F3F5')
    set_cell_border(cell,
        top={'val':'single','sz':'4','color':BORDER,'space':'0'},
        bottom={'val':'single','sz':'4','color':BORDER,'space':'0'},
        left={'val':'single','sz':'4','color':BORDER,'space':'0'},
        right={'val':'single','sz':'4','color':BORDER,'space':'0'})
    set_cell_margins(cell, 120, 150, 120, 150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(code.rstrip())
    set_run_font(run, 'Courier New', 7.2, NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Arial'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in [
        ('Title', 34, NAVY, 0, 6),
        ('Subtitle', 13, MUTED, 0, 8),
        ('Heading 1', 20, NAVY, 14, 7),
        ('Heading 2', 14, NAVY_2, 10, 5),
        ('Heading 3', 11, '765A22', 7, 3),
    ]:
        style = styles[name]
        style.font.name = 'Georgia' if name != 'Subtitle' else 'Arial'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), style.font.name)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != 'Subtitle'
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    for style_name in ['List Bullet', 'List Number']:
        st = styles[style_name]
        st.font.name = 'Arial'
        st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        st.font.size = Pt(8.9)
        st.paragraph_format.space_after = Pt(2)
        st.paragraph_format.left_indent = Cm(0.55)
        st.paragraph_format.first_line_indent = Cm(-0.22)


def add_cover(doc: Document):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0,0)
    set_cell_shading(cell, NAVY)
    set_cell_margins(cell, 700, 500, 700, 500)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('THEORETICAL CONFLICT ANALYSIS UNIT')
    set_run_font(r, 'Arial', 9, GOLD, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(15)
    r = p2.add_run('What Would Win')
    set_run_font(r, 'Georgia', 36, WHITE, bold=True)
    p3 = cell.add_paragraph()
    r = p3.add_run('Product planning, simulation specification and Codex handoff')
    set_run_font(r, 'Arial', 15, GOLD_LIGHT)
    p4 = cell.add_paragraph()
    p4.paragraph_format.space_before = Pt(18)
    r = p4.add_run('A transparent one-versus-X simulator for serious-looking answers to absurd questions.')
    set_run_font(r, 'Arial', 12, WHITE)
    doc.add_paragraph()

    info = doc.add_table(rows=4, cols=2)
    items = [
        ('Version', '0.2 — user-test candidate'),
        ('Prepared', '18 July 2026'),
        ('Owner / host', 'Samfa12-tech · samfa12.com'),
        ('Implementation', 'React · TypeScript · Vite · static hosting'),
    ]
    for idx, (label, value) in enumerate(items):
        left, right = info.rows[idx].cells
        set_cell_shading(left, PALE)
        set_cell_shading(right, 'FFFFFF')
        for cell in (left, right):
            set_cell_border(cell,
                top={'val':'single','sz':'4','color':BORDER,'space':'0'},
                bottom={'val':'single','sz':'4','color':BORDER,'space':'0'},
                left={'val':'single','sz':'4','color':BORDER,'space':'0'},
                right={'val':'single','sz':'4','color':BORDER,'space':'0'})
            set_cell_margins(cell, 100, 120, 100, 120)
        add_inline(left.paragraphs[0], label, base_size=8.5)
        left.paragraphs[0].runs[0].bold = True
        add_inline(right.paragraphs[0], value, base_size=8.5)
    doc.add_paragraph()
    add_callout(doc, 'Product promise: Put one creature against an effectively unlimited number of another creature, state the assumptions, and return a serious-looking, transparent, textual estimate of what would win.', fill='F6F2E7')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(32)
    r = p.add_run('Entertainment model · peak-adult profiles · no graphic violence')
    set_run_font(r, 'Arial', 9, MUTED, italic=True)
    doc.add_page_break()


def add_contents_map(doc: Document, headings: list[tuple[int, str]]):
    p = doc.add_paragraph('Document map', style='Heading 1')
    p.paragraph_format.space_before = Pt(0)
    add_styled_paragraph(doc, 'Use Word’s Navigation pane to jump between numbered sections. The map below lists the primary sections in this handoff.', spacing_after=8)
    for level, title in headings:
        if level != 2 or title == 'Product planning, simulation specification and Codex handoff':
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.45)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.paragraph_format.space_after = Pt(2.5)
        add_inline(p, f'•  {title}', base_size=8.9)
    doc.add_page_break()


def extract_headings(lines: Iterable[str]):
    result = []
    for line in lines:
        m = re.match(r'^(#{1,3})\s+(.+)$', line.strip())
        if m:
            result.append((len(m.group(1)), m.group(2)))
    return result


def build() -> None:
    lines = MARKDOWN.read_text(encoding='utf-8').splitlines()
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    section.header_distance = Cm(0.65)
    section.footer_distance = Cm(0.65)

    # Header/footer
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run('WHAT WOULD WIN · PRODUCT PLAN · v0.2')
    set_run_font(r, 'Arial', 7.5, MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    r = footer.add_run('samfa12.com · prototype handoff')
    set_run_font(r, 'Arial', 7.5, MUTED)
    footer.add_run('\t')
    add_page_number(footer)

    add_cover(doc)
    add_contents_map(doc, extract_headings(lines))

    # Skip markdown title and metadata already represented on cover.
    i = 0
    in_code = False
    code_lines: list[str] = []
    inserted_inputs = False
    inserted_results = False
    major_page_breaks: set[str] = set()

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()
        if i < 11 and (stripped.startswith('#') or stripped.startswith('**Document') or stripped.startswith('**Status') or stripped.startswith('**Prepared') or stripped.startswith('**Owner') or stripped.startswith('**Primary') or stripped == ''):
            i += 1
            continue
        if stripped == '---':
            i += 1
            continue
        if stripped.startswith('```'):
            if in_code:
                add_code_block(doc, '\n'.join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if stripped.startswith('|'):
            rows, i = parse_table(lines, i)
            add_markdown_table(doc, rows)
            continue
        heading = re.match(r'^(#{1,3})\s+(.+)$', stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if text in {'What Would Win', 'Product planning, simulation specification and Codex handoff'}:
                i += 1
                continue
            if text in major_page_breaks:
                doc.add_page_break()
            if text == '10.5 Structural integrity' and not inserted_results:
                doc.add_picture(str(ASSETS / 'demo-result.png'), width=Cm(16.7))
                cap = doc.add_paragraph('Figure 2. Working desktop prototype: textual verdict, metrics and transparent explanation.')
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cap.runs:
                    set_run_font(run, 'Arial', 7.5, MUTED, italic=True)
                inserted_results = True
            p = doc.add_paragraph(text, style=f'Heading {level}')
            if level == 1:
                p.paragraph_format.keep_with_next = True
            if text == '8. Information architecture and interface' and not inserted_inputs:
                doc.add_picture(str(ASSETS / 'demo-inputs.png'), width=Cm(16.7))
                cap = doc.add_paragraph('Figure 1. Working desktop prototype: contestant dossiers and primary controls.')
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cap.runs:
                    set_run_font(run, 'Arial', 7.5, MUTED, italic=True)
                inserted_inputs = True
            i += 1
            continue
        if stripped.startswith('> '):
            add_callout(doc, stripped[2:])
            i += 1
            continue
        bullet = re.match(r'^[-*]\s+(.+)$', stripped)
        numbered = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if bullet:
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, bullet.group(1), base_size=8.9)
            i += 1
            continue
        if numbered:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.55)
            p.paragraph_format.first_line_indent = Cm(-0.32)
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, f'{numbered.group(1)}. {numbered.group(2)}', base_size=8.9)
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        add_styled_paragraph(doc, stripped)
        i += 1

    # Core document properties
    props = doc.core_properties
    props.title = 'What Would Win — Product planning, simulation specification and Codex handoff'
    props.subject = 'Product requirements and implementation handoff'
    props.author = 'Samfa12-tech'
    props.keywords = 'What Would Win, simulation, React, TypeScript, product plan, Codex'
    props.comments = 'Public product planning document updated 18 July 2026.'

    doc.save(OUT)
    print(OUT)


if __name__ == '__main__':
    build()
